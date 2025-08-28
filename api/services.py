import json
import uuid
import asyncio
import logging
from typing import List, Dict, Any
from django.conf import settings
from openai import AsyncOpenAI
from core.models import ProcessNode, NodeDocument, NodeUsecaseCandidate, NodeEmbedding, AdminSettings

logger = logging.getLogger(__name__)


class OpenAIService:
    def __init__(self):
        # Try admin settings first, fallback to environment variables
        api_key = AdminSettings.get_setting('openai_api_key') or settings.OPENAI_API_KEY
        self.model = AdminSettings.get_setting('openai_model') or settings.OPENAI_MODEL or 'gpt-4o'
        
        # Log configuration details
        api_key_source = "Admin Settings" if AdminSettings.get_setting('openai_api_key') else "Environment Variables"
        model_source = "Admin Settings" if AdminSettings.get_setting('openai_model') else "Default/Environment"
        
        logger.info(f"🔑 OpenAI API Key source: {api_key_source}")
        logger.info(f"🔑 API Key present: {'Yes' if api_key else 'No'}")
        if api_key:
            logger.info(f"🔑 API Key preview: {api_key[:10]}...{api_key[-4:] if len(api_key) > 14 else '****'}")
        
        logger.info(f"🎯 OpenAI Model source: {model_source}")
        logger.info(f"🎯 Model: {self.model}")
        
        if not api_key:
            logger.error("🔑 ❌ No OpenAI API key found!")
            raise ValueError("OpenAI API key not found. Please set it in Admin Settings or environment variables.")
        
        self.client = AsyncOpenAI(api_key=api_key)
        logger.info("🤖 OpenAI AsyncClient initialized successfully")
    
    def generate_process_details(self, node: ProcessNode, context: Dict[str, Any]) -> Dict[str, Any]:
        """Generate detailed process description with AI"""
        try:
            return asyncio.run(self._async_generate_process_details(node, context))
        except Exception as e:
            logger.error(f"🤖 ❌ Failed to run async generate_process_details: {str(e)}")
            raise
    
    async def _async_generate_process_details(self, node: ProcessNode, context: Dict[str, Any]) -> Dict[str, Any]:
        """Async implementation of process details generation"""
        prompt = self._build_process_details_prompt(node, context)
        
        logger.info(f"🤖 [PROCESS_DETAILS] Starting OpenAI process details generation for node: {node.code}")
        logger.info(f"🤖 [PROCESS_DETAILS] Model: {self.model}")
        logger.info(f"🤖 [PROCESS_DETAILS] Prompt length: {len(prompt)} characters")
        logger.info(f"🤖 [PROCESS_DETAILS] Full prompt:\n{'-'*80}\n{prompt}\n{'-'*80}")
        
        system_message = ("You are a process modeling expert. Compose a clean, accurate, and practical process description. "
                         "Ground your answer in the provided context (ancestors, children, siblings/cousins, nearest neighbors, optional branch subtree). "
                         "If cross-category dependencies are plausible, include them explicitly with reasons. "
                         "Respond ONLY with valid JSON for the given schema. Keep it concise and business-focused. Return JSON only.")
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt}
        ]
        
        logger.info(f"🤖 [PROCESS_DETAILS] System message: {system_message}")
        logger.info(f"🤖 [PROCESS_DETAILS] Messages structure: {len(messages)} messages")
        
        # Making OpenAI call (no temperature parameter to avoid model restrictions)
        logger.info("🤖 [PROCESS_DETAILS] Making OpenAI call...")
        
        try:
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                response_format={"type": "json_object"}
                # No temperature - use model default
            )
            
            logger.info("🤖 [PROCESS_DETAILS] ✅ OpenAI call successful")
            
        except Exception as api_error:
            logger.error(f"🤖 [PROCESS_DETAILS] ❌ OpenAI call failed: {str(api_error)}")
            logger.error(f"🤖 [PROCESS_DETAILS] Exception type: {type(api_error).__name__}")
            raise api_error
        
        # Log response details
        logger.info(f"🤖 [PROCESS_DETAILS] Response received:")
        logger.info(f"🤖 [PROCESS_DETAILS] - Model used: {response.model}")
        logger.info(f"🤖 [PROCESS_DETAILS] - Total tokens: {response.usage.total_tokens}")
        logger.info(f"🤖 [PROCESS_DETAILS] - Prompt tokens: {response.usage.prompt_tokens}")
        logger.info(f"🤖 [PROCESS_DETAILS] - Completion tokens: {response.usage.completion_tokens}")
        logger.info(f"🤖 [PROCESS_DETAILS] - Finish reason: {response.choices[0].finish_reason}")
        
        content = response.choices[0].message.content
        logger.info(f"🤖 [PROCESS_DETAILS] Response content length: {len(content)} characters")
        logger.info(f"🤖 [PROCESS_DETAILS] Full response content:\n{'-'*80}\n{content}\n{'-'*80}")
        
        try:
            parsed_result = json.loads(content)
            logger.info(f"🤖 [PROCESS_DETAILS] ✅ JSON parsing successful. Keys: {list(parsed_result.keys())}")
            return parsed_result
        except json.JSONDecodeError as e:
            logger.error(f"🤖 [PROCESS_DETAILS] ❌ JSON parsing failed: {str(e)}")
            logger.error(f"🤖 [PROCESS_DETAILS] Raw content that failed to parse: {content}")
            raise e
    
    def generate_usecase_candidates(self, node: ProcessNode, context: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Generate AI use case candidates for a process"""
        return asyncio.run(self._async_generate_usecase_candidates(node, context))
    
    async def _async_generate_usecase_candidates(self, node: ProcessNode, context: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Async implementation of usecase candidates generation"""
        prompt = self._build_usecase_prompt(node, context)
        
        logger.info(f"🤖 [USECASE_CANDIDATES] Starting OpenAI use case generation for node: {node.code}")
        logger.info(f"🤖 [USECASE_CANDIDATES] Model: {self.model}")
        logger.info(f"🤖 [USECASE_CANDIDATES] Prompt length: {len(prompt)} characters")
        logger.info(f"🤖 [USECASE_CANDIDATES] Full prompt:\n{'-'*80}\n{prompt}\n{'-'*80}")
        
        system_message = "You are an AI strategy consultant. Generate practical AI use case candidates in JSON format."
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt}
        ]
        
        logger.info(f"🤖 [USECASE_CANDIDATES] System message: {system_message}")
        logger.info(f"🤖 [USECASE_CANDIDATES] Making OpenAI call...")
        
        try:
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                response_format={"type": "json_object"}
                # No temperature - use model default
            )
            
            logger.info(f"🤖 [USECASE_CANDIDATES] ✅ OpenAI call successful")
            logger.info(f"🤖 [USECASE_CANDIDATES] - Model used: {response.model}")
            logger.info(f"🤖 [USECASE_CANDIDATES] - Total tokens: {response.usage.total_tokens}")
            logger.info(f"🤖 [USECASE_CANDIDATES] - Prompt tokens: {response.usage.prompt_tokens}")
            logger.info(f"🤖 [USECASE_CANDIDATES] - Completion tokens: {response.usage.completion_tokens}")
            
            content = response.choices[0].message.content
            logger.info(f"🤖 [USECASE_CANDIDATES] Response content length: {len(content)} characters")
            logger.info(f"🤖 [USECASE_CANDIDATES] Full response:\n{'-'*80}\n{content}\n{'-'*80}")
            
            result = json.loads(content)
            logger.info(f"🤖 [USECASE_CANDIDATES] ✅ JSON parsing successful. Found {len(result.get('candidates', []))} candidates")
            return result.get('candidates', [])
            
        except Exception as e:
            logger.error(f"🤖 [USECASE_CANDIDATES] ❌ Failed: {str(e)}")
            raise
    
    def generate_usecase_specification(self, candidate: NodeUsecaseCandidate, context: Dict[str, Any]) -> str:
        """Generate detailed use case specification"""
        return asyncio.run(self._async_generate_usecase_specification(candidate, context))
    
    async def _async_generate_usecase_specification(self, candidate: NodeUsecaseCandidate, context: Dict[str, Any]) -> str:
        """Async implementation of usecase specification generation"""
        prompt = self._build_specification_prompt(candidate, context)
        
        logger.info(f"🤖 [USECASE_SPEC] Starting OpenAI specification generation for candidate: {candidate.title}")
        logger.info(f"🤖 [USECASE_SPEC] Model: {self.model}")
        logger.info(f"🤖 [USECASE_SPEC] Prompt length: {len(prompt)} characters")
        logger.info(f"🤖 [USECASE_SPEC] Full prompt:\n{'-'*80}\n{prompt}\n{'-'*80}")
        
        system_message = "You are a technical business analyst. Generate a comprehensive use case specification in markdown format."
        messages = [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt}
        ]
        
        logger.info(f"🤖 [USECASE_SPEC] System message: {system_message}")
        logger.info(f"🤖 [USECASE_SPEC] Making OpenAI call...")
        
        try:
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=messages
                # No temperature - use model default
            )
            
            logger.info(f"🤖 [USECASE_SPEC] ✅ OpenAI call successful")
            logger.info(f"🤖 [USECASE_SPEC] - Model used: {response.model}")
            logger.info(f"🤖 [USECASE_SPEC] - Total tokens: {response.usage.total_tokens}")
            logger.info(f"🤖 [USECASE_SPEC] - Prompt tokens: {response.usage.prompt_tokens}")
            logger.info(f"🤖 [USECASE_SPEC] - Completion tokens: {response.usage.completion_tokens}")
            
            content = response.choices[0].message.content
            logger.info(f"🤖 [USECASE_SPEC] Response content length: {len(content)} characters")
            logger.info(f"🤖 [USECASE_SPEC] Full response:\n{'-'*80}\n{content}\n{'-'*80}")
            
            return content
            
        except Exception as e:
            logger.error(f"🤖 [USECASE_SPEC] ❌ Failed: {str(e)}")
            raise
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for similarity search"""
        return asyncio.run(self._async_generate_embeddings(texts))
    
    async def _async_generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Async implementation of embeddings generation"""
        logger.info(f"🤖 [EMBEDDINGS] Starting OpenAI embeddings generation")
        logger.info(f"🤖 [EMBEDDINGS] Model: text-embedding-3-small")
        logger.info(f"🤖 [EMBEDDINGS] Number of texts: {len(texts)}")
        logger.info(f"🤖 [EMBEDDINGS] Total characters: {sum(len(t) for t in texts)}")
        
        for i, text in enumerate(texts[:3]):  # Log first 3 texts as sample
            preview = text[:200] + "..." if len(text) > 200 else text
            logger.info(f"🤖 [EMBEDDINGS] Sample text {i+1}: {preview}")
        
        try:
            response = await self.client.embeddings.create(
                model="text-embedding-3-small",
                input=texts
            )
            
            logger.info(f"🤖 [EMBEDDINGS] ✅ OpenAI embeddings call successful")
            logger.info(f"🤖 [EMBEDDINGS] - Model used: {response.model}")
            logger.info(f"🤖 [EMBEDDINGS] - Embeddings generated: {len(response.data)}")
            logger.info(f"🤖 [EMBEDDINGS] - Embedding dimensions: {len(response.data[0].embedding) if response.data else 0}")
            logger.info(f"🤖 [EMBEDDINGS] - Total tokens used: {response.usage.total_tokens}")
            
            return [embedding.embedding for embedding in response.data]
            
        except Exception as e:
            logger.error(f"🤖 [EMBEDDINGS] ❌ Failed: {str(e)}")
            raise
    
    def _build_process_details_prompt(self, node: ProcessNode, context: Dict[str, Any]) -> str:
        """Build comprehensive prompt matching old system's richness"""
        import json
        
        # Build the user request object matching old system
        user_request = {
            "task": "Compose a complete process description for the node.",
            "node": context.get("node"),
            "ancestors": context.get("ancestors", []),
            "children": context.get("children", []),
            "attributes": context.get("attributes", []),
            "candidates": {
                "siblings": context.get("siblings", []),
                "cousins": context.get("cousins", []),
                "nearest_neighbors": context.get("nearest_neighbors", [])
            },
            "guidance": {
                "allow_cross_category": context.get("cross_category", False),
                "note": "Consider cross-category flows (e.g., purchase -> warranty claim) when relevant."
            },
            "branch_context": context.get("branch_context", []),
            "schema": {
                "type": "object",
                "properties": {
                    "title": {"type": "string"},
                    "overview": {"type": "string"},
                    "inputs": {"type": "array", "items": {"type": "string"}},
                    "outputs": {"type": "array", "items": {"type": "string"}},
                    "kpis": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "description": {"type": "string"},
                                "formula": {"type": "string"},
                                "data_requirements": {"type": "array", "items": {"type": "string"}}
                            }
                        }
                    },
                    "upstream_processes": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "code": {"type": "string"},
                                "name": {"type": "string"},
                                "reason": {"type": "string"}
                            }
                        }
                    },
                    "downstream_processes": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "code": {"type": "string"},
                                "name": {"type": "string"},
                                "reason": {"type": "string"}
                            }
                        }
                    },
                    "related_processes": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "code": {"type": "string"},
                                "name": {"type": "string"},
                                "reason": {"type": "string"}
                            }
                        }
                    },
                    "steps": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "number": {"type": "integer"},
                                "name": {"type": "string"},
                                "description": {"type": "string"}
                            }
                        }
                    },
                    "challenges": {"type": "array", "items": {"type": "string"}},
                    "best_practices": {"type": "array", "items": {"type": "string"}}
                },
                "required": ["title", "overview", "inputs", "outputs"]
            }
        }
        
        # Convert to JSON string for the prompt
        return json.dumps(user_request, indent=2)
    
    def _build_usecase_prompt(self, node: ProcessNode, context: Dict[str, Any]) -> str:
        # Start with basic process information
        prompt_parts = [
            "Generate AI use case candidates for this business process:",
            "",
            f"Process Code: {node.code}",
            f"Process Name: {node.name}",
            f"Process Level: {node.level}",
            f"Description: {node.description or 'No description available'}",
            ""
        ]
        
        # Add process hierarchy context
        if context.get('ancestors'):
            prompt_parts.append("Process Hierarchy:")
            for ancestor in context['ancestors']:
                prompt_parts.append(f"  - Level {ancestor['level']}: {ancestor['code']} - {ancestor['name']}")
            prompt_parts.append("")
        
        # Add the full process details document if available
        if context.get('process_details'):
            prompt_parts.extend([
                "DETAILED PROCESS INFORMATION:",
                "================================",
                context['process_details'],
                "================================",
                ""
            ])
        
        # Add children processes if any
        if context.get('children'):
            prompt_parts.append("Sub-processes:")
            for child in context['children'][:10]:  # Limit to 10 to avoid token overflow
                prompt_parts.append(f"  - {child['code']}: {child['name']}")
            prompt_parts.append("")
        
        # Add sibling processes for context
        if context.get('siblings'):
            prompt_parts.append("Related Processes at Same Level:")
            for sibling in context['siblings'][:5]:  # Limit to 5
                prompt_parts.append(f"  - {sibling['code']}: {sibling['name']}")
            prompt_parts.append("")
        
        # Add the generation instructions
        prompt_parts.extend([
            "Based on the detailed process information above, generate 3-5 highly relevant and practical AI use case candidates.",
            "Consider the specific inputs, outputs, KPIs, steps, and challenges mentioned in the process details.",
            "Each use case should directly address the pain points and opportunities identified in the process documentation.",
            "",
            "Generate the candidates in JSON format:",
            "{",
            '    "candidates": [',
            "        {",
            '            "title": "Clear, actionable title that addresses a specific process need",',
            '            "description": "Detailed description explaining how AI will improve this specific process",',
            '            "impact_assessment": "Quantifiable business impact based on the process KPIs and objectives",',
            '            "complexity_score": 1-10,',
            '            "ai_technologies": ["specific AI technologies relevant to the process needs"],',
            '            "implementation_effort": "Low/Medium/High",',
            '            "roi_potential": "Low/Medium/High",',
            '            "process_alignment": "How this aligns with the process inputs, outputs, and KPIs"',
            "        }",
            "    ]",
            "}"
        ])
        
        prompt = "\n".join(prompt_parts)
        return prompt
    
    def _build_specification_prompt(self, candidate: NodeUsecaseCandidate, context: Dict[str, Any]) -> str:
        prompt = f"""
        Create a comprehensive technical specification for this AI use case:
        
        Use Case: {candidate.title}
        Process: {candidate.node.code} - {candidate.node.name}
        Description: {candidate.description}
        Impact Assessment: {candidate.impact_assessment}
        
        Generate a detailed markdown specification including:
        - Executive Summary
        - Business Requirements
        - Technical Requirements
        - Implementation Plan
        - Success Metrics
        - Risk Assessment
        - Resource Requirements
        """
        return prompt


class ContextService:
    @staticmethod
    def get_process_context(node: ProcessNode, include_branch: bool = False, cross_category: bool = False) -> Dict[str, Any]:
        """Get comprehensive contextual information for a process node (matching old system)"""
        logger.info(f"📊 [CONTEXT] Gathering context for node: {node.code}")
        logger.info(f"📊 [CONTEXT] Include branch: {include_branch}, Cross category: {cross_category}")
        
        context = {
            'node': {
                'id': node.id,
                'code': node.code,
                'name': node.name,
                'description': node.description,
                'level': node.level,
                'parent_id': node.parent_id if node.parent else None
            },
            'attributes': [],
            'ancestors': [],
            'children': [],
            'siblings': [],
            'cousins': [],
            'nearest_neighbors': [],
            'branch_context': []
        }
        
        # Get node attributes
        attributes = node.attributes.all().values('key', 'value')
        context['attributes'] = list(attributes)
        logger.info(f"📊 [CONTEXT] Found {len(context['attributes'])} attributes")
        
        # Get ancestors (full chain from root to node)
        ancestors = []
        current = node.parent
        while current:
            ancestors.append({
                'id': current.id,
                'code': current.code,
                'name': current.name,
                'level': current.level
            })
            current = current.parent
        context['ancestors'] = list(reversed(ancestors))  # Root to node order
        logger.info(f"📊 [CONTEXT] Found {len(context['ancestors'])} ancestors")
        
        # Get children
        children = node.children.all().values('code', 'name', 'description')
        context['children'] = list(children)
        logger.info(f"📊 [CONTEXT] Found {len(context['children'])} children")
        
        # Get siblings (same parent, excluding self)
        if node.parent:
            siblings = ProcessNode.objects.filter(
                parent=node.parent,
                model_version=node.model_version
            ).exclude(id=node.id).values('code', 'name', 'description')[:25]
            context['siblings'] = list(siblings)
            
            # Get cousins (children of parent's siblings)
            if node.parent.parent:  # Node has grandparent
                aunts_uncles = ProcessNode.objects.filter(
                    parent=node.parent.parent,
                    model_version=node.model_version
                ).exclude(id=node.parent_id)
                
                cousins = ProcessNode.objects.filter(
                    parent__in=aunts_uncles,
                    level=node.level,
                    model_version=node.model_version
                ).values('code', 'name', 'description')[:25]
                context['cousins'] = list(cousins)
        
        # Get nearest neighbors through embeddings (up to 60)
        if hasattr(node, 'embedding'):
            similar_nodes = ContextService._find_similar_nodes(node, limit=60)
            context['nearest_neighbors'] = similar_nodes
            
            # Diversify neighbors if cross_category is True
            if cross_category and similar_nodes:
                context['nearest_neighbors'] = ContextService._diversify_neighbors(similar_nodes, max_total=30)
        
        # Branch context (entire L2 subtree)
        if include_branch and node.level >= 2:
            l2_ancestor = ContextService._get_l2_ancestor(node)
            if l2_ancestor:
                branch_nodes = ProcessNode.objects.filter(
                    model_version=node.model_version,
                    materialized_path__startswith=l2_ancestor.materialized_path
                ).exclude(id=node.id).order_by('level', 'display_order').values('code', 'name', 'description', 'level')
                context['branch_context'] = list(branch_nodes)
        
        # Log summary
        logger.info(f"📊 [CONTEXT] Context summary: {len(context['siblings'])} siblings, "
                   f"{len(context['cousins'])} cousins, {len(context['nearest_neighbors'])} neighbors, "
                   f"{len(context['branch_context'])} branch nodes")
        
        return context
    
    @staticmethod
    def _find_similar_nodes(node: ProcessNode, limit: int = 10) -> List[Dict[str, Any]]:
        """Find similar nodes using embedding similarity"""
        # TODO: Implement proper cosine similarity search when embeddings are available
        # For now, return related nodes by parent relationship
        if node.parent:
            related = ProcessNode.objects.filter(
                parent=node.parent,
                model_version=node.model_version
            ).exclude(id=node.id).values('code', 'name', 'description')[:limit]
            return list(related)
        return []
    
    @staticmethod
    def _diversify_neighbors(neighbors: List[Dict], max_total: int = 30, max_per_prefix: int = 5) -> List[Dict]:
        """Diversify neighbors to include cross-category processes"""
        from collections import defaultdict
        
        # Group by L1 prefix (e.g., "1", "2", "3")
        grouped = defaultdict(list)
        for n in neighbors:
            prefix = n['code'].split('.')[0]
            grouped[prefix].append(n)
        
        # Take up to max_per_prefix from each group
        diversified = []
        for prefix in sorted(grouped.keys()):
            diversified.extend(grouped[prefix][:max_per_prefix])
            if len(diversified) >= max_total:
                break
        
        return diversified[:max_total]
    
    @staticmethod
    def _get_l2_ancestor(node: ProcessNode) -> ProcessNode:
        """Get the level 2 ancestor of a node"""
        current = node
        while current and current.level > 2:
            current = current.parent
        return current if current and current.level == 2 else None


class DocumentService:
    @staticmethod
    def save_document(node: ProcessNode, document_type: str, content: str, 
                     title: str = None, meta_data: Dict = None) -> NodeDocument:
        """Save a shared document"""
        document = NodeDocument.objects.create(
            node=node,
            document_type=document_type,
            title=title,
            content=content,
            meta_json=meta_data or {}
        )
        return document
    
    @staticmethod
    def find_document(node: ProcessNode, document_type: str) -> NodeDocument:
        """Find existing document for a node"""
        try:
            return NodeDocument.objects.get(
                node=node,
                document_type=document_type
            )
        except NodeDocument.DoesNotExist:
            return None
    
    @staticmethod
    def export_to_docx(document: NodeDocument) -> bytes:
        """Export document content to DOCX format"""
        from docx import Document
        from io import BytesIO
        
        doc = Document()
        doc.add_heading(document.title or f"{document.node.code} - {document.document_type}", 0)
        
        # Add content (assuming markdown)
        paragraphs = document.content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()